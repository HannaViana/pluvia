# -*- coding: utf-8 -*-
"""
Aplica correcoes ao Paper_20260227_mod.docx e salva Paper_20260227_mod_v2.docx.
Cada alteracao e marcada no texto com um tag entre colchetes.
Exclui a Secao 3.8 (sera tratada separadamente).
"""

from docx import Document
from lxml import etree
import sys

INPUT  = r"c:\Users\Hanna\work\projects\rain-and-flood-analysis\artigo\Paper_20260227_mod.docx"
OUTPUT = r"c:\Users\Hanna\work\projects\rain-and-flood-analysis\artigo\Paper_20260227_mod_v2.docx"

doc = Document(INPUT)

# -----------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------

def para_text(para):
    return "".join(r.text for r in para.runs)

def set_para_text(para, new_text):
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""

def replace_in_para(para, old, new):
    full = para_text(para)
    if old not in full:
        return False
    set_para_text(para, full.replace(old, new))
    return True

def replace_all(doc, old, new, tag=""):
    label = "  [" + tag + "]" if tag else ""
    count = 0
    for para in doc.paragraphs:
        if replace_in_para(para, old, new + label):
            count += 1
    return count

def delete_para(para):
    p = para._element
    p.getparent().remove(p)

# -----------------------------------------------------------------------
# C1 — §3.2: "blue curve therefore," -> "blue curve, therefore,"  [R1-C15]
# -----------------------------------------------------------------------
n = replace_all(doc,
    "The blue curve therefore, increases",
    "The blue curve, therefore, increases",
    "CORRECAO R1-C15")
print("[C1] blue curve therefore: %d substituicao(es)" % n)

# -----------------------------------------------------------------------
# C2 — §3.2: "All data" -> "All Data"  [R1-C12]
# -----------------------------------------------------------------------
n = replace_all(doc,
    "A separate box labeled All data (shown in the inset of panel b)",
    "A separate box labeled All Data (shown in the inset of panel b)",
    "CORRECAO R1-C12")
print("[C2a] All data label: %d substituicao(es)" % n)

n = replace_all(doc,
    "The All data summary box shows that when events longer than 12 hours are included",
    "The All Data inset shows that when events longer than 12 hours are included",
    "CORRECAO R1-C12")
print("[C2b] All data summary box: %d substituicao(es)" % n)

# -----------------------------------------------------------------------
# C3 — Introducao: adicionar [15] apos "how risk is perceived"
# Tenta variantes do texto para contornar diferenca de espaco/run
# -----------------------------------------------------------------------
candidates_c3 = [
    "how risk is perceived. This perspective helps clarify",
    "how risk is perceived.  This perspective helps clarify",
]
n_c3 = 0
for cand in candidates_c3:
    n_c3 += replace_all(doc, cand,
        "how risk is perceived [15]. This perspective helps clarify",
        "INSERCAO ref.[15]")
print("[C3] referencia [15]: %d substituicao(es)" % n_c3)

# -----------------------------------------------------------------------
# C4 — §3.6: corrigir numeracao das figuras nos paineis b e c
# -----------------------------------------------------------------------
n = replace_all(doc,
    "In autumn (Figure 10b), the hotspot pattern",
    "In autumn (Figure 9b), the hotspot pattern",
    "CORRECAO numeracao Fig.")
print("[C4a] Figure 10b->9b (autumn): %d" % n)

n = replace_all(doc,
    "In winter (Figure 10c), the density field contracts",
    "In winter (Figure 9c), the density field contracts",
    "CORRECAO numeracao Fig.")
print("[C4b] Figure 10c->9c (winter): %d" % n)

# -----------------------------------------------------------------------
# C5 — §3.6 inverno: adicionar fatores costeiros/mare  [tidal]
# -----------------------------------------------------------------------
n = replace_all(doc,
    ("which is consistent with chronic local vulnerability associated with "
     "drainage constraints and topographic depressions."),
    ("which is consistent with chronic local vulnerability associated with "
     "drainage constraints, topographic depressions, and, in some coastal areas, "
     "tidal boundary effects that can cause surface water to persist even after "
     "light rainfall [39]."),
    "INSERCAO tidal §3.6")
print("[C5] tidal factors winter §3.6: %d" % n)

# -----------------------------------------------------------------------
# C6 — §3.7: Thiessen = agora Fig.10
# -----------------------------------------------------------------------
fixes_thiessen = [
    ("Figure 9 combines two pieces of information in a single map",
     "Figure 10 combines two pieces of information in a single map"),
    ("Figure 9 shows strong spatial contrasts in the number of flood occurrences",
     "Figure 10 shows strong spatial contrasts in the number of flood occurrences"),
    ("because the allocation is based on Thiessen polygons, Figure 8 provides an operationally useful link",
     "because the allocation is based on Thiessen polygons, Figure 10 provides an operationally useful link"),
    ("From a threshold development perspective, Figure 9 is a key intermediate product",
     "From a threshold development perspective, Figure 10 is a key intermediate product"),
    ("Figure 9. Thiessen polygons for the Alerta Rio rain gauge network",
     "Figure 10. Thiessen polygons for the Alerta Rio rain gauge network"),
]
for old, new in fixes_thiessen:
    n = replace_all(doc, old, new, "CORRECAO numeracao Fig.")
    print("[C6] Thiessen Fig.9->10: %d  (%s)" % (n, old[:55]))

# -----------------------------------------------------------------------
# C7 — §3.7: distancia = agora Fig.11; sazonal espacial = Fig.9
# -----------------------------------------------------------------------
fixes_dist = [
    ("Figure 10 quantifies the distance between each georeferenced occurrence",
     "Figure 11 quantifies the distance between each georeferenced occurrence"),
    ("Figure 10 shows that the rain gauge network generally covers the occurrence",
     "Figure 11 shows that the rain gauge network generally covers the occurrence"),
    ("Figure 10. Distance from each occurrence to the nearest Alerta Rio",
     "Figure 11. Distance from each occurrence to the nearest Alerta Rio"),
    ("The seasonal spatial evolution in Figure 10 is consistent",
     "The seasonal spatial evolution in Figure 9 is consistent"),
    ("Importantly, Figure 10 adds a spatial dimension",
     "Importantly, Figure 9 adds a spatial dimension"),
]
for old, new in fixes_dist:
    n = replace_all(doc, old, new, "CORRECAO numeracao Fig.")
    print("[C7] Distance Fig.10->11/9: %d  (%s)" % (n, old[:55]))

# -----------------------------------------------------------------------
# C8 — Referencias: [36] (Georganta) -> [30]
# -----------------------------------------------------------------------
n = replace_all(doc,
    "Georganta et al. [36] defined intensity-duration thresholds",
    "Georganta et al. [30] defined intensity-duration thresholds",
    "CORRECAO ref.[36]->[30]")
print("[C8a] Georganta [36]->[30] intro: %d" % n)

n = replace_all(doc,
    "conducted in the Attica region of Greece [36]",
    "conducted in the Attica region of Greece [30]",
    "CORRECAO ref.[36]->[30]")
print("[C8b] Georganta [36]->[30] §2.2: %d" % n)

# -----------------------------------------------------------------------
# C9 — Renumerar referencias [37]->[36], [38]->[37], [39]->[38], [40]->[39]
# Usa marcadores temporarios para evitar conflitos.
# -----------------------------------------------------------------------
temp_map  = {"[37]": "[ZREF36]", "[38]": "[ZREF37]", "[39]": "[ZREF38]", "[40]": "[ZREF39]"}
final_map = {"[ZREF36]": "[36]",  "[ZREF37]": "[37]",  "[ZREF38]": "[38]",  "[ZREF39]": "[39]"}

for old, tmp in temp_map.items():
    cnt = 0
    for para in doc.paragraphs:
        if replace_in_para(para, old, tmp):
            cnt += 1
    print("[C9-temp] %s->%s: %d" % (old, tmp, cnt))

for tmp, new in final_map.items():
    cnt = 0
    for para in doc.paragraphs:
        if replace_in_para(para, tmp, new):
            cnt += 1
    print("[C9-final] %s->%s: %d" % (tmp, new, cnt))

# Remover linha da referencia [36] duplicata de Georganta
to_delete = []
for para in doc.paragraphs:
    txt = para_text(para)
    if txt.strip().startswith("[36] Georganta"):
        to_delete.append(para)
        print("[C9-delref] Marcado para remocao: %s" % txt[:80])
for para in to_delete:
    delete_para(para)
    print("[C9-delref] Paragrafo removido.")

# -----------------------------------------------------------------------
# C10 — §3.5: remover bloco duplicado (texto antigo com "urban corridors")
# O bloco comeca com MARKER_START e termina com MARKER_END (inclusive).
# -----------------------------------------------------------------------
MARKER_START = ("Figure 8 maps the spatial concentration of occurrence records "
                "using density surfaces for (a) all types combined and (b to d) "
                "each occurrence type separately. The colored patches represent "
                "relative density classes of occurrence concentration, categorized "
                "from Very Low to Very High")

MARKER_END = ("A practical methodological note follows directly from the way "
              "Figure 8 is constructed.")

inside_block = False
to_delete = []

for para in doc.paragraphs:
    txt = para_text(para)
    if not inside_block and MARKER_START in txt:
        inside_block = True
        to_delete.append(para)
        print("[C10] Inicio bloco duplicado: %s" % txt[:80])
        continue
    if inside_block:
        to_delete.append(para)
        if MARKER_END in txt:
            print("[C10] Fim bloco duplicado: %s" % txt[:80])
            inside_block = False
            break

print("[C10] Paragrafos para remover: %d" % len(to_delete))
for para in to_delete:
    delete_para(para)
print("[C10] Bloco duplicado removido.")

# -----------------------------------------------------------------------
# Salvar
# -----------------------------------------------------------------------
doc.save(OUTPUT)
print("\nArquivo salvo: %s" % OUTPUT)
